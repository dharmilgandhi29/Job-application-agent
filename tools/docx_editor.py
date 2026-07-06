"""
docx_editor.py — the in-place tailor.

Edits your REAL Word resume by swapping the TEXT of specific paragraphs while
leaving every bit of formatting — fonts, tab-stopped dates, hyperlinks, bullet
styles, spacing, one-page layout — exactly as you designed it. It never rebuilds
the document; it finds a paragraph by its current text and rewrites just that text.
"""

from docx import Document
from docx.oxml.ns import qn


def _replace_paragraph_text(paragraph, new_text: str):
    """Replace a paragraph's entire text, preserving formatting, by writing
    new_text into the first run and clearing the rest. Avoids the multi-run trap."""
    if not paragraph.runs:
        paragraph.text = new_text
        return
    paragraph.runs[0].text = new_text
    for r in paragraph.runs[1:]:
        r.text = ""


def _replace_skill_line(paragraph, new_full_line: str):
    """Skills lines are 'Bold Category: normal, skill, list'. We split at the first
    colon: label+colon keep the first run's bold; the skill list goes into a second
    run with bold OFF but the SAME FONT as the original (copied explicitly, since a
    fresh run would otherwise default to a different font and break the layout)."""
    if ":" not in new_full_line or not paragraph.runs:
        _replace_paragraph_text(paragraph, new_full_line)
        return

    label, after = new_full_line.split(":", 1)
    first = paragraph.runs[0]

    # Capture the original font BEFORE editing, so the new run can match it exactly
    font_name = first.font.name
    font_size = first.font.size

    # First run: the bold label (keeps its existing bold formatting)
    first.text = f"{label}:"
    for r in paragraph.runs[1:]:
        r.text = ""

    # New run for the skill list: bold OFF, same font as the original
    new_run = paragraph.add_run(after)
    new_run.bold = False
    new_run.font.name = font_name
    new_run.font.size = font_size
    # Force the font at the XML level too — setting font.name alone doesn't always
    # "stick" on a freshly added run (known python-docx quirk).
    if font_name:
        rpr = new_run._element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn('w:ascii'), font_name)
        rfonts.set(qn('w:hAnsi'), font_name)


def apply_swaps(in_path: str, out_path: str, swaps: list[dict]) -> dict:
    """Open in_path, apply swaps, save to out_path.

    Each swap: {"original": "<exact current paragraph text>",
                "new": "<new text>",
                "kind": "bullet" | "summary" | "skill"}
    Matches by stripped original text. Returns {applied, missed} for verification."""
    doc = Document(in_path)
    applied, missed = [], []

    for swap in swaps:
        original = swap["original"].strip()
        new_text = swap["new"]
        kind = swap.get("kind", "bullet")
        found = False
        for p in doc.paragraphs:
            if p.text.strip() == original:
                # Skills need the bold-label / normal-list split (with font copy);
                # everything else is a straight first-run text replacement.
                if kind == "skill":
                    _replace_skill_line(p, new_text)
                else:
                    _replace_paragraph_text(p, new_text)
                applied.append(original[:50])
                found = True
                break
        if not found:
            missed.append(original[:50])

    doc.save(out_path)
    return {"applied": applied, "missed": missed}


# ── Reading the swappable paragraphs OUT of the docx (fallback; the LLM classifier
#    in classify_resume.py is the primary path) ─────────────────────────────────
def extract_swappable(in_path: str) -> list[dict]:
    """Heuristic extraction of tailorable paragraphs. Kept as a fallback; the
    Claude-based classifier in classify_resume.py is what the pipeline actually uses."""
    doc = Document(in_path)
    out = []
    section = None

    for p in doc.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        up = t.upper()

        if up.startswith("PROFESSIONAL SUMMARY"):
            section = "summary"; continue
        if up.startswith("EDUCATION"):
            section = "education"; continue
        if up.startswith("WORK EXPERIENCE"):
            section = "work"; continue
        if up.startswith("ACADEMIC PROJECTS"):
            section = "projects"; continue
        if up.startswith("TECHNICAL SKILLS"):
            section = "skills"; continue

        if section == "summary":
            out.append({"original": t, "kind": "summary"})
            section = None
        elif section in ("work", "projects"):
            if "\t" in p.text:
                continue
            if " | " in t:
                continue
            if len(t) < 40 and t.endswith(("Intern", "Analyst", "Engineer")):
                continue
            if len(t) > 60:
                out.append({"original": t, "kind": "bullet"})
        elif section == "skills":
            if ":" in t:
                out.append({"original": t, "kind": "skill"})

    return out